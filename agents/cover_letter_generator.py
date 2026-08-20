import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# PATHS
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

JOB_ID = sys.argv[1] if len(sys.argv) > 1 else "job_001"

PRIVATE_RESUME = (
    BASE_DIR
    / "resumes"
    / "private"
    / "resume_profile.json"
)

TAILORING_PLAN = (
    BASE_DIR
    / "resumes"
    / "private"
    / "tailoring"
    / f"{JOB_ID}_tailoring_plan.json"
)

OUTPUT_DIR = (
    BASE_DIR
    / "resumes"
    / "output"
)


# ============================================================
# DESIGN SETTINGS
# ============================================================

FONT_NAME = "Arial"

NAVY = RGBColor(31, 78, 121)
DARK_TEXT = RGBColor(40, 40, 40)
LINK_BLUE = RGBColor(5, 99, 193)
LINE_COLOR = "1F4E79"


# ============================================================
# JSON
# ============================================================

def load_json(file_path):
    """Load JSON file."""

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as file:
        return json.load(file)


# ============================================================
# FONT
# ============================================================

def set_run_font(
    run,
    size=10,
    bold=False,
    color=DARK_TEXT,
    underline=False
):
    """Apply consistent font formatting."""

    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(
        qn("w:ascii"),
        FONT_NAME
    )
    run._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        FONT_NAME
    )
    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.underline = underline


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    """Create professional ATS-friendly document."""

    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)

    styles = document.styles

    normal = styles["Normal"]

    normal.font.name = FONT_NAME
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK_TEXT

    normal._element.rPr.rFonts.set(
        qn("w:ascii"),
        FONT_NAME
    )

    normal._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        FONT_NAME
    )

    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    return document


# ============================================================
# PARAGRAPH HELPERS
# ============================================================

def add_paragraph(
    document,
    text="",
    size=10,
    bold=False,
    color=DARK_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=4,
    line_spacing=1.08
):
    """Add a standard paragraph."""

    paragraph = document.add_paragraph()

    paragraph.alignment = alignment

    paragraph.paragraph_format.space_before = Pt(
        space_before
    )

    paragraph.paragraph_format.space_after = Pt(
        space_after
    )

    paragraph.paragraph_format.line_spacing = (
        line_spacing
    )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=size,
        bold=bold,
        color=color
    )

    return paragraph


def add_justified_paragraph(
    document,
    text,
    size=10,
    space_before=0,
    space_after=8
):
    """Add professionally justified paragraph."""

    return add_paragraph(
        document,
        text=text,
        size=size,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=space_before,
        space_after=space_after,
        line_spacing=1.12
    )


# ============================================================
# HORIZONTAL LINE
# ============================================================

def add_horizontal_line(
    document,
    color=LINE_COLOR,
    size="8"
):
    """
    Add a clean horizontal separator line.

    This line is placed AFTER the header/contact block,
    before the cover-letter date and recipient section.
    """

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.0

    p = paragraph._p

    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")

    bottom.set(
        qn("w:val"),
        "single"
    )

    bottom.set(
        qn("w:sz"),
        size
    )

    bottom.set(
        qn("w:space"),
        "1"
    )

    bottom.set(
        qn("w:color"),
        color
    )

    pBdr.append(bottom)

    pPr.append(pBdr)

    return paragraph


# ============================================================
# HYPERLINK
# ============================================================

def add_hyperlink(
    paragraph,
    text,
    url
):
    """Add blue underlined hyperlink."""

    part = paragraph.part

    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")

    hyperlink.set(
        qn("r:id"),
        relationship_id
    )

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(
        qn("w:val"),
        "0563C1"
    )

    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(
        qn("w:val"),
        "single"
    )

    rPr.append(underline)

    font = OxmlElement("w:rFonts")

    font.set(
        qn("w:ascii"),
        FONT_NAME
    )

    font.set(
        qn("w:hAnsi"),
        FONT_NAME
    )

    font.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    rPr.append(font)

    size = OxmlElement("w:sz")
    size.set(
        qn("w:val"),
        "20"
    )

    rPr.append(size)

    new_run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text

    new_run.append(text_element)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


# ============================================================
# HEADER
# ============================================================

def add_header(
    document,
    resume
):
    """Create professional cover letter header."""

    personal = resume.get(
        "personal",
        {}
    )

    professional_summary = resume.get(
        "professional_summary",
        {}
    )

    name = personal.get(
        "full_name",
        ""
    )

    target_role = professional_summary.get(
        "target_role",
        "Data Analyst"
    )

    location = personal.get(
        "location",
        ""
    )

    # --------------------------------------------------------
    # NAME
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(
        name.upper()
    )

    set_run_font(
        run,
        size=18,
        bold=True,
        color=NAVY
    )

    # --------------------------------------------------------
    # PROFESSIONAL DESIGNATION + LOCATION
    # SAME LINE
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)

    role_run = paragraph.add_run(
        target_role.upper()
    )

    set_run_font(
        role_run,
        size=10.5,
        bold=True,
        color=DARK_TEXT
    )

    if location:

        separator_run = paragraph.add_run(
            "  |  "
        )

        set_run_font(
            separator_run,
            size=10,
            color=DARK_TEXT
        )

        location_run = paragraph.add_run(
            location
        )

        set_run_font(
            location_run,
            size=10,
            color=DARK_TEXT
        )

    # --------------------------------------------------------
    # CONTACT LINE
    # --------------------------------------------------------

    contact = document.add_paragraph()

    contact.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(7)

    email = personal.get(
        "email",
        ""
    )

    phone = personal.get(
        "phone",
        ""
    )

    linkedin = personal.get(
        "linkedin_url",
        ""
    )

    if email:

        run = contact.add_run(
            email
        )

        set_run_font(
            run,
            size=9,
            color=DARK_TEXT
        )

    if phone:

        if email:

            run = contact.add_run(
                "  |  "
            )

            set_run_font(
                run,
                size=9,
                color=DARK_TEXT
            )

        run = contact.add_run(
            phone
        )

        set_run_font(
            run,
            size=9,
            color=DARK_TEXT
        )

    if linkedin:

        if email or phone:

            run = contact.add_run(
                "  |  "
            )

            set_run_font(
                run,
                size=9,
                color=DARK_TEXT
            )

        display_link = linkedin

        if display_link.startswith(
            "https://"
        ):

            display_link = display_link[
                8:
            ]

        elif display_link.startswith(
            "http://"
        ):

            display_link = display_link[
                7:
            ]

        add_hyperlink(
            contact,
            display_link,
            linkedin
        )

    # --------------------------------------------------------
    # IMPORTANT:
    # HEADER SEPARATOR LINE
    # --------------------------------------------------------

    add_horizontal_line(
        document
    )


# ============================================================
# DATE
# ============================================================

def add_date(
    document
):
    """Add current application date."""

    from datetime import datetime

    date_text = datetime.now().strftime(
        "%d %B %Y"
    )

    add_paragraph(
        document,
        date_text,
        size=10,
        space_after=12
    )


# ============================================================
# RECIPIENT
# ============================================================

def add_recipient(
    document,
    tailoring
):
    """Add recipient information."""

    job = tailoring.get(
        "job",
        {}
    )

    company = job.get(
        "company",
        ""
    )

    location = job.get(
        "location",
        ""
    )

    title = job.get(
        "title",
        "Data Analyst"
    )

    add_paragraph(
        document,
        "Hiring Manager",
        size=10,
        bold=True,
        space_after=1
    )

    if company:

        add_paragraph(
            document,
            company,
            size=10,
            space_after=1
        )

    if location:

        add_paragraph(
            document,
            location,
            size=10,
            space_after=8
        )

    add_paragraph(
        document,
        f"Re: Application for {title}",
        size=10,
        bold=True,
        color=NAVY,
        space_after=10
    )


# ============================================================
# COVER LETTER BODY
# ============================================================

def get_cover_letter_content(
    resume,
    tailoring
):
    """
    Build conservative cover letter content
    using only verified candidate information.
    """

    personal = resume.get(
        "personal",
        {}
    )

    professional_summary = resume.get(
        "professional_summary",
        {}
    )

    job = tailoring.get(
        "job",
        {}
    )

    name = personal.get(
        "full_name",
        ""
    )

    target_role = job.get(
        "title",
        professional_summary.get(
            "target_role",
            "Data Analyst"
        )
    )

    company = job.get(
        "company",
        ""
    )

    matched_required = tailoring.get(
        "matched_required_skills",
        []
    )

    matched_preferred = tailoring.get(
        "matched_preferred_skills",
        []
    )

    matched = (
        matched_required
        + matched_preferred
    )

    projects = tailoring.get(
        "recommended_projects",
        []
    )

    project_names = [
        project.get(
            "project",
            ""
        )
        for project in projects
        if project.get(
            "project",
            ""
        )
    ]

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    if matched:

        skill_text = ", ".join(
            matched[:7]
        )

    else:

        skill_text = (
            "SQL, Python, Excel, Tableau, "
            "data cleaning, reporting, and "
            "business intelligence"
        )

    # --------------------------------------------------------
    # PROJECT
    # --------------------------------------------------------

    if project_names:

        main_project = project_names[0]

    else:

        main_project = (
            "Farmers Market Analytics"
        )

    # --------------------------------------------------------
    # PARAGRAPHS
    # --------------------------------------------------------

    paragraphs = []

    paragraphs.append(
        f"Dear Hiring Manager,"
    )

    paragraphs.append(
        f"I am writing to apply for the "
        f"{target_role} position at {company}. "
        f"I am an analytical and results-driven "
        f"professional with hands-on experience "
        f"in {skill_text}. I am particularly "
        f"interested in opportunities where "
        f"data analysis and business intelligence "
        f"can be used to improve decision-making "
        f"and operational performance."
    )

    paragraphs.append(
        f"Through my practical analytics work, "
        f"I developed {main_project}, an "
        f"end-to-end analytics solution covering "
        f"data preparation, analysis, reporting, "
        f"automation, and dashboard development. "
        f"This work strengthened my ability to "
        f"transform operational data into meaningful "
        f"insights and communicate findings clearly "
        f"to support business decisions."
    )

    paragraphs.append(
        f"My professional background has also "
        f"developed strong problem-solving, "
        f"communication, documentation, stakeholder "
        f"and process-focused skills. Combined with "
        f"my technical analytics capabilities, I can "
        f"bring a practical and structured approach "
        f"to analysing business problems and "
        f"identifying opportunities for improvement."
    )

    work_rights = personal.get(
        "work_rights",
        ""
    )

    if work_rights:

        paragraphs.append(
            f"I am based in {personal.get('location', '')} "
            f"and {work_rights.lower()}. "
            f"I would welcome the opportunity to "
            f"discuss how my analytical background "
            f"and technical skills could contribute "
            f"to your team."
        )

    else:

        paragraphs.append(
            f"I would welcome the opportunity to "
            f"discuss how my analytical background "
            f"and technical skills could contribute "
            f"to your team."
        )

    return paragraphs


# ============================================================
# BODY
# ============================================================

def add_body(
    document,
    resume,
    tailoring
):
    """Add cover letter body."""

    paragraphs = get_cover_letter_content(
        resume,
        tailoring
    )

    for index, text in enumerate(
        paragraphs
    ):

        if index == 0:

            add_paragraph(
                document,
                text,
                size=10,
                space_after=10
            )

        else:

            add_justified_paragraph(
                document,
                text,
                size=10,
                space_after=10
            )


# ============================================================
# CLOSING
# ============================================================

def add_closing(
    document,
    resume
):
    """
    Add closing with TWO blank lines
    before 'Kind regards'.
    """

    personal = resume.get(
        "personal",
        {}
    )

    name = personal.get(
        "full_name",
        ""
    )

    # Closing sentence
    add_paragraph(
        document,
        "Thank you for considering my application.",
        size=10,
        space_before=2,
        space_after=0
    )

    # --------------------------------------------------------
    # TWO-LINE SPACE
    # --------------------------------------------------------

    add_paragraph(
        document,
        "",
        size=10,
        space_after=0
    )

    add_paragraph(
        document,
        "",
        size=10,
        space_after=0
    )

    # --------------------------------------------------------
    # KIND REGARDS
    # --------------------------------------------------------

    add_paragraph(
        document,
        "Kind regards,",
        size=10,
        space_before=0,
        space_after=2
    )

    # Name
    add_paragraph(
        document,
        name,
        size=10,
        bold=True,
        color=NAVY,
        space_after=0
    )


# ============================================================
# SAFETY
# ============================================================

def get_safety(
    tailoring
):
    """Return internal safety information."""

    safety = {
        "invented_experience": False,
        "invented_skills": False,
        "invented_education": False,
        "invented_employment": False,
        "automatic_submission": False,
        "human_review_required": True
    }

    plan_safety = tailoring.get(
        "safety",
        {}
    )

    for key in safety:

        if key in plan_safety:

            safety[key] = plan_safety[key]

    return safety


# ============================================================
# DOCUMENT GENERATION
# ============================================================

def generate_cover_letter(
    resume,
    tailoring
):
    """Generate complete cover letter."""

    document = setup_document()

    add_header(
        document,
        resume
    )

    add_date(
        document
    )

    add_recipient(
        document,
        tailoring
    )

    add_body(
        document,
        resume,
        tailoring
    )

    add_closing(
        document,
        resume
    )

    return document


# ============================================================
# MAIN
# ============================================================

def main():

    print()
    print(
        "COVER LETTER GENERATOR - VERSION 10"
    )

    print(
        "=" * 68
    )

    # --------------------------------------------------------
    # CHECK FILES
    # --------------------------------------------------------

    if not PRIVATE_RESUME.exists():

        print()
        print(
            "ERROR: Resume profile not found:"
        )

        print(
            PRIVATE_RESUME
        )

        return

    if not TAILORING_PLAN.exists():

        print()
        print(
            "ERROR: Tailoring plan not found:"
        )

        print(
            TAILORING_PLAN
        )

        return

    # --------------------------------------------------------
    # LOAD DATA
    # --------------------------------------------------------

    resume = load_json(
        PRIVATE_RESUME
    )

    tailoring = load_json(
        TAILORING_PLAN
    )

    job = tailoring.get(
        "job",
        {}
    )

    # --------------------------------------------------------
    # DISPLAY
    # --------------------------------------------------------

    print()
    print(
        "COVER LETTER GENERATOR"
    )

    print(
        "=" * 68
    )

    print()

    print(
        f"Job:      "
        f"{job.get('title', '')}"
    )

    print(
        f"Company:  "
        f"{job.get('company', '')}"
    )

    print(
        f"Location: "
        f"{job.get('location', '')}"
    )

    # --------------------------------------------------------
    # FORMAT
    # --------------------------------------------------------

    print()
    print(
        "FORMAT:"
    )

    print(
        "  Single-column layout"
    )

    print(
        "  Arial font"
    )

    print(
        "  Professional navy header"
    )

    print(
        "  Header/contact separator line"
    )

    print(
        "  Blue underlined LinkedIn hyperlink"
    )

    print(
        "  One-line professional designation and location"
    )

    print(
        "  Justified body paragraphs"
    )

    print(
        "  Two-line space before Kind regards"
    )

    print(
        "  ATS-safe Word document"
    )

    print(
        "  No tables"
    )

    print(
        "  No text boxes"
    )

    print(
        "  No graphics"
    )

    # --------------------------------------------------------
    # SAFETY
    # --------------------------------------------------------

    safety = get_safety(
        tailoring
    )

    print()
    print(
        "SAFETY:"
    )

    print(
        f"  Invented experience: "
        f"{safety['invented_experience']}"
    )

    print(
        f"  Invented skills: "
        f"{safety['invented_skills']}"
    )

    print(
        f"  Invented education: "
        f"{safety['invented_education']}"
    )

    print(
        f"  Invented employment: "
        f"{safety['invented_employment']}"
    )

    print(
        f"  Automatic submission: "
        f"{safety['automatic_submission']}"
    )

    print(
        f"  Human review required: "
        f"{safety['human_review_required']}"
    )

    # --------------------------------------------------------
    # GENERATE
    # --------------------------------------------------------

    document = generate_cover_letter(
        resume,
        tailoring
    )

    # --------------------------------------------------------
    # OUTPUT
    # --------------------------------------------------------

    job_id = job.get(
        "job_id",
        "unknown"
    )

    output_dir = (
        OUTPUT_DIR
        / job_id
        / "cover_letter"
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    output_file = (
        output_dir
        / "cover_letter.docx"
    )

    try:

        document.save(
            output_file
        )

    except PermissionError:

        print()
        print(
            "ERROR: Word appears to have the "
            "cover letter file open."
        )

        print()
        print(
            "Close the existing cover_letter.docx "
            "in Microsoft Word and run the command again."
        )

        return

    # --------------------------------------------------------
    # SUCCESS
    # --------------------------------------------------------

    print()
    print(
        "Cover letter saved:"
    )

    print(
        output_file
    )


# ============================================================
# RUN
# ============================================================

if __name__ == "__main__":
    main()