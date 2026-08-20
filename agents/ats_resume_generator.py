import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# VERSION 8
# ATS-FRIENDLY RESUME GENERATOR
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

JOB_ID = sys.argv[1] if len(sys.argv) > 1 else "job_001"

PRIVATE_RESUME = (
    BASE_DIR
    / "resumes"
    / "private"
    / "resume_profile.json"
)

TAILORING_PLAN = (
    BASE_DIR
    / "resumes"
    / "private"
    / "tailoring"
    / f"{JOB_ID}_tailoring_plan.json"
)

OUTPUT_DIR = (
    BASE_DIR
    / "resumes"
    / "output"
)


# ============================================================
# DESIGN SETTINGS
# ============================================================

FONT_NAME = "Arial"

NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(5, 99, 193)
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(64, 64, 64)

BODY_SIZE = 10
SMALL_SIZE = 9
NAME_SIZE = 18
DESIGNATION_SIZE = 10.5
SECTION_SIZE = 11.5
SUBHEADING_SIZE = 10.5

LINE_SPACING = 1.12


# ============================================================
# JSON HELPERS
# ============================================================

def load_json(file_path):
    """Load JSON data."""

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as file:
        return json.load(file)


# ============================================================
# FONT HELPERS
# ============================================================

def set_run_font(
    run,
    size=BODY_SIZE,
    bold=False,
    color=BLACK,
    italic=False,
    underline=False
):
    """Apply consistent font settings to a run."""

    run.font.name = FONT_NAME

    run._element.rPr.rFonts.set(
        qn("w:ascii"),
        FONT_NAME
    )

    run._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        FONT_NAME
    )

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = color


def set_paragraph_spacing(
    paragraph,
    before=0,
    after=4,
    line_spacing=LINE_SPACING
):
    """Apply consistent paragraph spacing."""

    paragraph.paragraph_format.space_before = Pt(
        before
    )

    paragraph.paragraph_format.space_after = Pt(
        after
    )

    paragraph.paragraph_format.line_spacing = (
        line_spacing
    )


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    """Create the ATS-friendly Word document."""

    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    styles = document.styles

    normal = styles["Normal"]

    normal.font.name = FONT_NAME

    normal._element.rPr.rFonts.set(
        qn("w:ascii"),
        FONT_NAME
    )

    normal._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        FONT_NAME
    )

    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )

    normal.font.size = Pt(BODY_SIZE)

    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = LINE_SPACING

    return document


# ============================================================
# DIVIDER LINE
# ============================================================

def add_top_border(paragraph):
    """
    Add a thin divider line ABOVE the section heading.

    This keeps the line visually attached to the section,
    rather than placing it below the heading.
    """

    p = paragraph._p

    pPr = p.get_or_add_pPr()

    pBdr = pPr.find(
        qn("w:pBdr")
    )

    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    top = OxmlElement("w:top")

    top.set(
        qn("w:val"),
        "single"
    )

    top.set(
        qn("w:sz"),
        "8"
    )

    top.set(
        qn("w:space"),
        "5"
    )

    top.set(
        qn("w:color"),
        "1F4E79"
    )

    pBdr.append(top)


# ============================================================
# HYPERLINK
# ============================================================

def add_hyperlink(
    paragraph,
    text,
    url
):
    """
    Add ATS-readable hyperlink.

    The visible URL remains text so ATS systems can
    still read the LinkedIn/GitHub address.
    """

    if not url:
        return

    part = paragraph.part

    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")

    hyperlink.set(
        qn("r:id"),
        relationship_id
    )

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(
        qn("w:val"),
        "0563C1"
    )
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(
        qn("w:val"),
        "single"
    )
    rPr.append(underline)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(
        qn("w:ascii"),
        FONT_NAME
    )
    rFonts.set(
        qn("w:hAnsi"),
        FONT_NAME
    )
    rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME
    )
    rPr.append(rFonts)

    size = OxmlElement("w:sz")
    size.set(
        qn("w:val"),
        "18"
    )
    rPr.append(size)

    new_run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text

    new_run.append(text_element)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


# ============================================================
# STANDARD TEXT
# ============================================================

def add_text(
    document,
    text,
    bold=False,
    size=BODY_SIZE,
    color=BLACK,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    before=0,
    after=4,
    line_spacing=LINE_SPACING
):
    """Add a normal ATS-safe paragraph."""

    paragraph = document.add_paragraph()

    paragraph.alignment = alignment

    set_paragraph_spacing(
        paragraph,
        before=before,
        after=after,
        line_spacing=line_spacing
    )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=size,
        bold=bold,
        color=color
    )

    return paragraph


# ============================================================
# JUSTIFIED PARAGRAPH
# ============================================================

def add_justified_paragraph(
    document,
    text,
    size=BODY_SIZE,
    before=0,
    after=7
):
    """Add a professionally justified paragraph."""

    if not text:
        return None

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    set_paragraph_spacing(
        paragraph,
        before=before,
        after=after,
        line_spacing=1.15
    )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=size,
        color=BLACK
    )

    return paragraph


# ============================================================
# SECTION HEADING
# ============================================================

def add_section_heading(
    document,
    text
):
    """
    Add section divider ABOVE heading.

    Layout:

    ------------------------------
    PROFESSIONAL SUMMARY

    Description...
    """

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0

    add_top_border(
        paragraph
    )

    run = paragraph.add_run(
        text.upper()
    )

    set_run_font(
        run,
        size=SECTION_SIZE,
        bold=True,
        color=NAVY
    )

    return paragraph


# ============================================================
# SUBHEADING
# ============================================================

def add_subheading(
    document,
    text,
    after=4
):
    """Add professional subheading."""

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    set_paragraph_spacing(
        paragraph,
        before=2,
        after=after,
        line_spacing=1.0
    )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=SUBHEADING_SIZE,
        bold=True,
        color=NAVY
    )

    return paragraph


# ============================================================
# BULLET
# ============================================================

def add_bullet(
    document,
    text
):
    """Add ATS-safe bullet with comfortable spacing."""

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = (
        Inches(0.18)
    )

    paragraph.paragraph_format.first_line_indent = (
        Inches(-0.12)
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12

    run = paragraph.add_run(
        "\u2022 "
        + str(text)
    )

    set_run_font(
        run,
        size=BODY_SIZE,
        color=BLACK
    )

    return paragraph


# ============================================================
# LABEL + VALUE
# ============================================================

def add_label_value(
    document,
    label,
    value
):
    """Add bold label followed by normal text."""

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.12

    label_run = paragraph.add_run(
        str(label)
    )

    set_run_font(
        label_run,
        size=BODY_SIZE,
        bold=True,
        color=NAVY
    )

    value_run = paragraph.add_run(
        str(value)
    )

    set_run_font(
        value_run,
        size=BODY_SIZE,
        color=BLACK
    )

    return paragraph


# ============================================================
# PERSONAL HEADER
# ============================================================

def add_personal_header(
    document,
    resume
):
    """
    Header layout:

    SARAVANAKUMAR KANNAN

    DATA ANALYST | AUCKLAND, NEW ZEALAND

    Email | Phone | LinkedIn | GitHub
    """

    personal = resume.get(
        "personal",
        {}
    )

    professional = resume.get(
        "professional_summary",
        {}
    )

    name = personal.get(
        "full_name",
        ""
    )

    target_role = professional.get(
        "target_role",
        "Data Analyst"
    )

    location = personal.get(
        "location",
        ""
    )

    # --------------------------------------------------------
    # NAME
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(
        name
    )

    set_run_font(
        run,
        size=NAME_SIZE,
        bold=True,
        color=NAVY
    )

    # --------------------------------------------------------
    # DESIGNATION + LOCATION SAME LINE
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0

    designation_run = paragraph.add_run(
        target_role.upper()
    )

    set_run_font(
        designation_run,
        size=DESIGNATION_SIZE,
        bold=True,
        color=BLACK
    )

    if location:

        separator_run = paragraph.add_run(
            "  |  "
        )

        set_run_font(
            separator_run,
            size=DESIGNATION_SIZE,
            color=DARK_GRAY
        )

        location_run = paragraph.add_run(
            location
        )

        set_run_font(
            location_run,
            size=DESIGNATION_SIZE,
            color=DARK_GRAY
        )

    # --------------------------------------------------------
    # CONTACT INFORMATION
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0

    contact_items = []

    email = personal.get(
        "email",
        ""
    )

    phone = personal.get(
        "phone",
        ""
    )

    linkedin = personal.get(
        "linkedin_url",
        ""
    )

    github = personal.get(
        "github_url",
        ""
    )

    if email:

        run = paragraph.add_run(
            "Email: "
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            bold=True,
            color=DARK_GRAY
        )

        run = paragraph.add_run(
            email
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            color=BLACK
        )

        contact_items.append(
            "email"
        )

    if phone:

        if contact_items:

            run = paragraph.add_run(
                "  |  "
            )

            set_run_font(
                run,
                size=SMALL_SIZE,
                color=DARK_GRAY
            )

        run = paragraph.add_run(
            "Phone: "
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            bold=True,
            color=DARK_GRAY
        )

        run = paragraph.add_run(
            phone
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            color=BLACK
        )

        contact_items.append(
            "phone"
        )

    if linkedin:

        if contact_items:

            run = paragraph.add_run(
                "  |  "
            )

            set_run_font(
                run,
                size=SMALL_SIZE,
                color=DARK_GRAY
            )

        run = paragraph.add_run(
            "LinkedIn: "
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            bold=True,
            color=DARK_GRAY
        )

        add_hyperlink(
            paragraph,
            linkedin,
            linkedin
        )

        contact_items.append(
            "linkedin"
        )

    if github:

        if contact_items:

            run = paragraph.add_run(
                "  |  "
            )

            set_run_font(
                run,
                size=SMALL_SIZE,
                color=DARK_GRAY
            )

        run = paragraph.add_run(
            "GitHub: "
        )

        set_run_font(
            run,
            size=SMALL_SIZE,
            bold=True,
            color=DARK_GRAY
        )

        add_hyperlink(
            paragraph,
            github,
            github
        )


# ============================================================
# PROFESSIONAL SUMMARY
# ============================================================

def build_professional_summary(
    resume,
    tailoring
):
    """Build verified targeted summary."""

    summary_data = resume.get(
        "professional_summary",
        {}
    )

    base_summary = summary_data.get(
        "summary",
        ""
    ).strip()

    matched_required = tailoring.get(
        "matched_required_skills",
        []
    )

    matched_preferred = tailoring.get(
        "matched_preferred_skills",
        []
    )

    matched = (
        matched_required
        + matched_preferred
    )

    if not base_summary and matched:

        return (
            "Analytical and results-driven "
            "Data Analyst with hands-on "
            "experience applying "
            + ", ".join(matched)
            + " through practical analytics "
              "projects."
        )

    if base_summary and matched:

        return (
            base_summary
            + " Experienced in applying "
            + ", ".join(matched)
            + " through practical analytics "
              "projects."
        )

    return base_summary


def add_professional_summary(
    document,
    resume,
    tailoring
):
    """Add professional summary."""

    add_section_heading(
        document,
        "Professional Summary"
    )

    summary = build_professional_summary(
        resume,
        tailoring
    )

    if summary:

        add_justified_paragraph(
            document,
            summary,
            size=BODY_SIZE,
            after=8
        )


# ============================================================
# CORE SKILLS
# ============================================================

def add_core_skills(
    document,
    resume,
    tailoring
):
    """Add targeted ATS keyword skills."""

    add_section_heading(
        document,
        "Core Skills"
    )

    skills = resume.get(
        "technical_skills",
        {}
    )

    candidate_skills = set()

    for group in skills.values():

        for skill in group:

            candidate_skills.add(
                skill.lower()
            )

    categories = {
        "Business Analysis": [
            "Business Analysis",
            "Process Optimisation",
            "Continuous Improvement",
            "KPI Development",
            "Performance Monitoring",
            "Trend Analysis",
            "Root Cause Analysis",
            "Business Reporting"
        ],
        "Data Analytics": [
            "SQL",
            "MySQL",
            "Python",
            "Pandas",
            "NumPy",
            "Data Cleaning",
            "Data Validation",
            "Data Transformation",
            "Data Automation"
        ],
        "Business Intelligence": [
            "Tableau",
            "Excel",
            "Advanced Excel",
            "PivotTables",
            "Dashboards",
            "KPI Reporting",
            "Data Visualisation",
            "Performance Analytics"
        ],
        "Professional Skills": [
            "Stakeholder Management",
            "Cross-functional Collaboration",
            "Problem Solving",
            "Communication",
            "Attention to Detail",
            "Project Coordination"
        ]
    }

    priority_skills = (
        tailoring.get(
            "priority_skills",
            []
        )
    )

    priority_lookup = {
        skill.lower()
        for skill in priority_skills
    }

    for category, category_skills in categories.items():

        verified = []

        for skill in category_skills:

            if skill.lower() in candidate_skills:
                verified.append(skill)

        # Put job-matched skills first.
        verified.sort(
            key=lambda skill: (
                0
                if skill.lower() in priority_lookup
                else 1
            )
        )

        if verified:

            add_label_value(
                document,
                category + ": ",
                " • ".join(verified)
            )


# ============================================================
# EXPERIENCE
# ============================================================

def add_experience(
    document,
    resume
):
    """Add verified professional experience."""

    experience = resume.get(
        "experience",
        []
    )

    valid_experience = [
        item
        for item in experience
        if item.get("company")
        or item.get("job_title")
    ]

    if not valid_experience:
        return

    add_section_heading(
        document,
        "Professional Experience"
    )

    for index, role in enumerate(
        valid_experience
    ):

        title = role.get(
            "job_title",
            ""
        )

        company = role.get(
            "company",
            ""
        )

        location = role.get(
            "location",
            ""
        )

        start = role.get(
            "start_date",
            ""
        )

        end = role.get(
            "end_date",
            ""
        )

        current = role.get(
            "current",
            False
        )

        if current:
            date_text = (
                start
                + "–Present"
                if start
                else "Present"
            )

        elif start and end:

            date_text = (
                start
                + "–"
                + end
            )

        else:

            date_text = (
                start
                or end
            )

        # ----------------------------------------------------
        # JOB TITLE
        # ----------------------------------------------------

        header_parts = []

        if title:
            header_parts.append(
                title
            )

        if company:
            header_parts.append(
                company
            )

        header_text = " | ".join(
            header_parts
        )

        add_subheading(
            document,
            header_text,
            after=2
        )

        # ----------------------------------------------------
        # LOCATION + DATES
        # ----------------------------------------------------

        metadata = []

        if location:
            metadata.append(
                location
            )

        if date_text:
            metadata.append(
                date_text
            )

        if metadata:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run(
                " | ".join(metadata)
            )

            set_run_font(
                run,
                size=SMALL_SIZE,
                italic=True,
                color=DARK_GRAY
            )

        # ----------------------------------------------------
        # RESPONSIBILITIES
        # ----------------------------------------------------

        for responsibility in role.get(
            "responsibilities",
            []
        ):

            if responsibility:

                add_bullet(
                    document,
                    responsibility
                )

        # ----------------------------------------------------
        # ACHIEVEMENTS
        # ----------------------------------------------------

        for achievement in role.get(
            "achievements",
            []
        ):

            if achievement:

                add_bullet(
                    document,
                    achievement
                )

        if index < len(valid_experience) - 1:

            spacer = document.add_paragraph()

            spacer.paragraph_format.space_after = Pt(2)


# ============================================================
# PROJECT HELPERS
# ============================================================

def get_project_map(resume):
    """Create project lookup."""

    project_map = {}

    for project in resume.get(
        "portfolio_projects",
        []
    ):

        name = project.get(
            "name",
            ""
        )

        if name:

            project_map[name] = project

    return project_map


def select_projects(
    resume,
    tailoring
):
    """Select strongest verified projects."""

    project_map = get_project_map(
        resume
    )

    recommendations = tailoring.get(
        "recommended_projects",
        []
    )

    selected = []

    for recommendation in recommendations:

        project_name = recommendation.get(
            "project",
            ""
        )

        project = project_map.get(
            project_name
        )

        if project:

            selected.append(
                (
                    project,
                    recommendation
                )
            )

    if selected:

        return selected

    return [
        (
            project,
            {
                "relevance_score": 0,
                "matching_skills": []
            }
        )
        for project in resume.get(
            "portfolio_projects",
            []
        )[:3]
    ]


# ============================================================
# PROJECTS
# ============================================================

def add_projects(
    document,
    resume,
    tailoring
):
    """Add selected analytics projects."""

    selected_projects = select_projects(
        resume,
        tailoring
    )

    if not selected_projects:
        return

    add_section_heading(
        document,
        "Selected Data Analytics Projects"
    )

    for project, recommendation in selected_projects:

        name = project.get(
            "name",
            ""
        )

        technologies = project.get(
            "technologies",
            []
        )

        title = name

        if technologies:

            title += (
                " | "
                + " | ".join(
                    technologies
                )
            )

        add_subheading(
            document,
            title,
            after=5
        )

        # ----------------------------------------------------
        # Project description / achievements
        # ----------------------------------------------------

        achievements = project.get(
            "achievements",
            []
        )

        for achievement in achievements:

            if achievement:

                add_bullet(
                    document,
                    achievement
                )

        # ----------------------------------------------------
        # Matching evidence
        # ----------------------------------------------------

        matching_skills = recommendation.get(
            "matching_skills",
            []
        )

        if matching_skills:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0

            label_run = paragraph.add_run(
                "Relevant skills: "
            )

            set_run_font(
                label_run,
                size=SMALL_SIZE,
                bold=True,
                color=NAVY
            )

            value_run = paragraph.add_run(
                ", ".join(
                    matching_skills
                )
            )

            set_run_font(
                value_run,
                size=SMALL_SIZE,
                color=DARK_GRAY
            )


# ============================================================
# EDUCATION
# ============================================================

def add_education(
    document,
    resume
):
    """Add education."""

    education = resume.get(
        "education",
        []
    )

    valid_education = [
        item
        for item in education
        if item.get("qualification")
    ]

    if not valid_education:
        return

    add_section_heading(
        document,
        "Education"
    )

    for item in valid_education:

        qualification = item.get(
            "qualification",
            ""
        )

        specialization = item.get(
            "specialization",
            ""
        )

        institution = item.get(
            "institution",
            ""
        )

        status = item.get(
            "status",
            ""
        )

        line = qualification

        if specialization:

            line += (
                " — "
                + specialization
            )

        add_subheading(
            document,
            line,
            after=2
        )

        metadata = []

        if institution:
            metadata.append(
                institution
            )

        if status:
            metadata.append(
                status
            )

        if metadata:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run(
                " | ".join(metadata)
            )

            set_run_font(
                run,
                size=SMALL_SIZE,
                color=DARK_GRAY
            )


# ============================================================
# TECHNICAL SKILLS
# ============================================================

def add_technical_skills(
    document,
    resume,
    tailoring
):
    """Add verified technical skills."""

    skills = resume.get(
        "technical_skills",
        {}
    )

    if not skills:
        return

    add_section_heading(
        document,
        "Technical Skills"
    )

    labels = {
        "programming": "Programming & Analytics",
        "databases": "Databases",
        "visualization": "Visualisation",
        "analytics": "Business Intelligence & Analytics",
        "machine_learning": "Machine Learning",
        "automation": "Automation",
        "development": "Development",
        "ai": "Artificial Intelligence"
    }

    priority_skills = tailoring.get(
        "priority_skills",
        []
    )

    priority_lookup = {
        skill.lower()
        for skill in priority_skills
    }

    for category, items in skills.items():

        if not items:
            continue

        ordered_items = list(items)

        ordered_items.sort(
            key=lambda skill: (
                0
                if skill.lower()
                in priority_lookup
                else 1
            )
        )

        label = labels.get(
            category,
            category.replace(
                "_",
                " "
            ).title()
        )

        add_label_value(
            document,
            label + ": ",
            " • ".join(
                ordered_items
            )
        )


# ============================================================
# CERTIFICATIONS
# ============================================================

def add_certifications(
    document,
    resume
):
    """Add certifications."""

    certifications = resume.get(
        "certifications",
        []
    )

    if not certifications:
        return

    add_section_heading(
        document,
        "Certifications"
    )

    add_text(
        document,
        " • ".join(
            certifications
        ),
        size=BODY_SIZE,
        after=6
    )


# ============================================================
# ADDITIONAL INFORMATION
# ============================================================

def add_additional_information(
    document,
    resume
):
    """Add verified additional information."""

    information = resume.get(
        "additional_information",
        {}
    )

    personal = resume.get(
        "personal",
        {}
    )

    items = []

    work_rights = personal.get(
        "work_rights",
        ""
    )

    if work_rights:

        items.append(
            work_rights
        )

    driving_license = information.get(
        "driving_license",
        ""
    )

    if driving_license:

        items.append(
            driving_license
        )

    availability = information.get(
        "availability",
        ""
    )

    if availability:

        items.append(
            availability
        )

    preferred_work_modes = information.get(
        "preferred_work_modes",
        []
    )

    if preferred_work_modes:

        items.append(
            "Preferred work modes: "
            + ", ".join(
                preferred_work_modes
            )
        )

    interests = information.get(
        "interests",
        []
    )

    if interests:

        items.append(
            "Professional interests: "
            + ", ".join(
                interests
            )
        )

    if not items:
        return

    add_section_heading(
        document,
        "Additional Information"
    )

    for item in items:

        add_bullet(
            document,
            item
        )


# ============================================================
# SAFETY
# ============================================================

def get_safety(
    tailoring
):
    """Return internal safety controls."""

    safety = {
        "invented_experience": False,
        "invented_skills": False,
        "invented_education": False,
        "invented_employment": False,
        "personal_data_exposed": False,
        "automatic_submission": False,
        "human_review_required": True
    }

    plan_safety = tailoring.get(
        "safety",
        {}
    )

    for key in safety:

        if key in plan_safety:

            safety[key] = (
                plan_safety[key]
            )

    return safety


# ============================================================
# RESUME GENERATION
# ============================================================

def generate_resume(
    resume,
    tailoring
):
    """Generate complete ATS-friendly resume."""

    document = setup_document()

    # Header
    add_personal_header(
        document,
        resume
    )

    # Main sections
    add_professional_summary(
        document,
        resume,
        tailoring
    )

    add_core_skills(
        document,
        resume,
        tailoring
    )

    add_experience(
        document,
        resume
    )

    add_projects(
        document,
        resume,
        tailoring
    )

    add_education(
        document,
        resume
    )

    add_technical_skills(
        document,
        resume,
        tailoring
    )

    add_certifications(
        document,
        resume
    )

    add_additional_information(
        document,
        resume
    )

    return document


# ============================================================
# MAIN
# ============================================================

def main():

    print()
    print(
        "ATS RESUME GENERATOR - VERSION 8"
    )

    print("=" * 68)

    # --------------------------------------------------------
    # Validate files
    # --------------------------------------------------------

    if not PRIVATE_RESUME.exists():

        print()
        print(
            "ERROR: Resume profile not found:"
        )

        print(
            PRIVATE_RESUME
        )

        return

    if not TAILORING_PLAN.exists():

        print()
        print(
            "ERROR: Tailoring plan not found:"
        )

        print(
            TAILORING_PLAN
        )

        return

    # --------------------------------------------------------
    # Load data
    # --------------------------------------------------------

    resume = load_json(
        PRIVATE_RESUME
    )

    tailoring = load_json(
        TAILORING_PLAN
    )

    job = tailoring.get(
        "job",
        {}
    )

    # --------------------------------------------------------
    # Display information
    # --------------------------------------------------------

    print()
    print(
        "ATS RESUME GENERATOR"
    )

    print("=" * 68)

    print()
    print(
        f"Job:      "
        f"{job.get('title', '')}"
    )

    print(
        f"Company:  "
        f"{job.get('company', '')}"
    )

    print(
        f"Location: "
        f"{job.get('location', '')}"
    )

    # --------------------------------------------------------
    # ATS format
    # --------------------------------------------------------

    print()
    print(
        "ATS FORMAT:"
    )

    print(
        "  Single-column layout"
    )

    print(
        "  Arial font"
    )

    print(
        "  Professional navy headings"
    )

    print(
        "  Blue underlined hyperlinks"
    )

    print(
        "  Section divider above each section"
    )

    print(
        "  Designation and location on same line"
    )

    print(
        "  Justified professional paragraphs"
    )

    print(
        "  Increased line spacing"
    )

    print(
        "  Increased subheading spacing"
    )

    print(
        "  Standard section headings"
    )

    print(
        "  No tables"
    )

    print(
        "  No text boxes"
    )

    print(
        "  No graphics"
    )

    print(
        "  No icons"
    )

    print(
        "  Standard Word document"
    )

    # --------------------------------------------------------
    # Safety
    # --------------------------------------------------------

    safety = get_safety(
        tailoring
    )

    print()
    print(
        "SAFETY:"
    )

    print(
        f"  Invented experience: "
        f"{safety['invented_experience']}"
    )

    print(
        f"  Invented skills: "
        f"{safety['invented_skills']}"
    )

    print(
        f"  Invented education: "
        f"{safety['invented_education']}"
    )

    print(
        f"  Invented employment: "
        f"{safety['invented_employment']}"
    )

    print(
        f"  Automatic submission: "
        f"{safety['automatic_submission']}"
    )

    print(
        f"  Human review required: "
        f"{safety['human_review_required']}"
    )

    # --------------------------------------------------------
    # Generate
    # --------------------------------------------------------

    document = generate_resume(
        resume,
        tailoring
    )

    # --------------------------------------------------------
    # Output directory
    # --------------------------------------------------------

    job_id = job.get(
        "job_id",
        "unknown"
    )

    job_output_dir = (
        OUTPUT_DIR
        / job_id
        / "tailored"
    )

    job_output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    output_file = (
        job_output_dir
        / "resume.docx"
    )

    # --------------------------------------------------------
    # Save
    # --------------------------------------------------------

    try:

        document.save(
            output_file
        )

    except PermissionError:

        print()
        print(
            "ERROR: Word has the previous resume "
            "file open."
        )

        print()
        print(
            "Close resume.docx in Microsoft Word "
            "and run the command again."
        )

        return

    # --------------------------------------------------------
    # Complete
    # --------------------------------------------------------

    print()
    print(
        "Resume saved:"
    )

    print(
        output_file
    )

    print()
    print(
        "VERSION 8 COMPLETE"
    )


if __name__ == "__main__":
    main()